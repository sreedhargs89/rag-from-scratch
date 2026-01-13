"""
Module 1: Document Loading

This module teaches you how to load and extract text from various document formats.
Understanding this is crucial for RAG because you need clean, structured text before
you can chunk and embed it.

Key Concepts:
- Reading different file formats (TXT, PDF, DOCX, MD)
- Extracting clean text while preserving structure
- Handling metadata (filename, page numbers, etc.)
"""

import os
from pathlib import Path
from typing import List, Dict, Optional
import json


class Document:
    """
    Represents a loaded document with its content and metadata.

    This is the fundamental data structure we'll use throughout the RAG pipeline.
    """
    def __init__(self, content: str, metadata: Optional[Dict] = None):
        self.content = content
        self.metadata = metadata or {}

    def __repr__(self):
        content_preview = self.content[:100] + "..." if len(self.content) > 100 else self.content
        return f"Document(content='{content_preview}', metadata={self.metadata})"

    def to_dict(self):
        """Convert to dictionary for serialization"""
        return {
            "content": self.content,
            "metadata": self.metadata
        }


class DocumentLoader:
    """
    Loads documents from various file formats.

    This class demonstrates the first step in the RAG pipeline: getting text
    from files into a format we can work with.
    """

    def __init__(self):
        self.supported_extensions = {'.txt', '.md', '.pdf', '.docx', '.json'}

    def load_file(self, file_path: str) -> Document:
        """
        Load a single file and return a Document object.

        Args:
            file_path: Path to the file

        Returns:
            Document object containing the text and metadata
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower()

        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {extension}")

        # Extract metadata
        metadata = {
            "source": str(path),
            "filename": path.name,
            "extension": extension,
            "size_bytes": path.stat().st_size
        }

        # Load content based on file type
        if extension == '.txt' or extension == '.md':
            content = self._load_text(path)
        elif extension == '.pdf':
            content = self._load_pdf(path)
        elif extension == '.docx':
            content = self._load_docx(path)
        elif extension == '.json':
            content = self._load_json(path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

        return Document(content=content, metadata=metadata)

    def load_directory(self, dir_path: str, recursive: bool = True) -> List[Document]:
        """
        Load all supported documents from a directory.

        Args:
            dir_path: Path to directory
            recursive: Whether to search subdirectories

        Returns:
            List of Document objects
        """
        path = Path(dir_path)

        if not path.exists():
            raise FileNotFoundError(f"Directory not found: {dir_path}")

        documents = []

        if recursive:
            files = path.rglob("*")
        else:
            files = path.glob("*")

        for file_path in files:
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    doc = self.load_file(str(file_path))
                    documents.append(doc)
                    print(f"✓ Loaded: {file_path.name}")
                except Exception as e:
                    print(f"✗ Error loading {file_path.name}: {e}")

        return documents

    def _load_text(self, path: Path) -> str:
        """Load plain text or markdown files"""
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()

    def _load_pdf(self, path: Path) -> str:
        """
        Load PDF files using pdfplumber.

        Note: Install with: pip install pdfplumber
        """
        try:
            import pdfplumber
        except ImportError:
            return "ERROR: pdfplumber not installed. Install with: pip install pdfplumber"

        text_parts = []

        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    # Add page marker for reference
                    text_parts.append(f"[Page {page_num}]\n{text}")

        return "\n\n".join(text_parts)

    def _load_docx(self, path: Path) -> str:
        """
        Load DOCX files using python-docx.

        Note: Install with: pip install python-docx
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            return "ERROR: python-docx not installed. Install with: pip install python-docx"

        doc = DocxDocument(path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n\n".join(paragraphs)

    def _load_json(self, path: Path) -> str:
        """
        Load JSON files.

        This converts JSON to readable text format.
        """
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Convert to formatted string
        return json.dumps(data, indent=2)


# ============================================================================
# EXAMPLES AND EXERCISES
# ============================================================================

def example_load_single_file():
    """Example: Load a single file"""
    print("=" * 60)
    print("Example 1: Loading a Single File")
    print("=" * 60)

    loader = DocumentLoader()

    # Create a sample file
    sample_file = "../data/sample_docs/sample.txt"
    os.makedirs(os.path.dirname(sample_file), exist_ok=True)

    with open(sample_file, 'w') as f:
        f.write("""Machine Learning Basics

Machine learning is a subset of artificial intelligence that enables systems to learn
and improve from experience without being explicitly programmed.

There are three main types of machine learning:
1. Supervised Learning - Learning from labeled data
2. Unsupervised Learning - Finding patterns in unlabeled data
3. Reinforcement Learning - Learning through trial and error

Applications include image recognition, natural language processing, and recommendation systems.
""")

    # Load the document
    doc = loader.load_file(sample_file)

    print(f"\nDocument loaded!")
    print(f"Source: {doc.metadata['source']}")
    print(f"Size: {doc.metadata['size_bytes']} bytes")
    print(f"\nContent preview:\n{doc.content[:200]}...")


def example_load_directory():
    """Example: Load all files from a directory"""
    print("\n" + "=" * 60)
    print("Example 2: Loading Multiple Files")
    print("=" * 60)

    loader = DocumentLoader()

    # Create sample files
    sample_dir = "../data/sample_docs/"
    os.makedirs(sample_dir, exist_ok=True)

    # Create a few sample documents
    samples = {
        "ai_overview.txt": "Artificial Intelligence is the simulation of human intelligence by machines.",
        "deep_learning.txt": "Deep learning uses neural networks with multiple layers to learn hierarchical representations.",
        "nlp_intro.md": "# Natural Language Processing\n\nNLP enables computers to understand and generate human language."
    }

    for filename, content in samples.items():
        with open(os.path.join(sample_dir, filename), 'w') as f:
            f.write(content)

    # Load all documents
    documents = loader.load_directory(sample_dir, recursive=False)

    print(f"\nLoaded {len(documents)} documents:")
    for doc in documents:
        print(f"  - {doc.metadata['filename']}: {len(doc.content)} characters")


def exercise_1():
    """
    EXERCISE 1: Inspect Document Metadata

    Task: Load a document and print all its metadata fields.
    """
    print("\n" + "=" * 60)
    print("EXERCISE 1: Inspect Document Metadata")
    print("=" * 60)

    # TODO: Load a document and print its metadata
    # Hint: Use loader.load_file() and access doc.metadata

    print("\nYour code here!")


def exercise_2():
    """
    EXERCISE 2: Filter Documents by Size

    Task: Load multiple documents and filter out those smaller than 100 characters.
    """
    print("\n" + "=" * 60)
    print("EXERCISE 2: Filter Documents by Size")
    print("=" * 60)

    # TODO: Load documents from a directory and filter by size

    print("\nYour code here!")


if __name__ == "__main__":
    print("\n" + "🚀 RAG FROM SCRATCH - MODULE 1: DOCUMENT LOADING" + "\n")

    # Run examples
    example_load_single_file()
    example_load_directory()

    # Exercises
    exercise_1()
    exercise_2()

    print("\n" + "=" * 60)
    print("Key Takeaways:")
    print("=" * 60)
    print("✓ Documents are the starting point for RAG")
    print("✓ We need clean text extraction from various formats")
    print("✓ Metadata helps track source and context")
    print("✓ Next: We'll learn how to split these documents into chunks")
    print("=" * 60)
